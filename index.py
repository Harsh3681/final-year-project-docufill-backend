from flask import Flask, Response,request, send_file
import time
from flask_cors import CORS
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig, TextIteratorStreamer
import torch
from threading import Thread
from docx import Document
import os
# from docx2pdf import convert
from PIL import Image
from io import BytesIO
from base64 import b64decode
import requests
import gc
from dotenv import load_dotenv
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import convertapi

load_dotenv()
HF_TOKEN = os.getenv('HF_TOKEN')
convertapi.api_secret = os.getenv("API_SECRET")


sections = {
    "Title":0,
    "Abstract":7,
    "Introduction":10,
    "ProposedMethod":14,
    "RelatedWork":14,
    "Experiments":16,
    "Conclusion":18,
}

bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16
)

model = AutoModelForCausalLM.from_pretrained("Shrunoti09/llama2",quantization_config=bnb_config,
    device_map={"": 0},
    low_cpu_mem_usage=True,
    trust_remote_code=True)

model.config.use_cache = False

tokenizer = AutoTokenizer.from_pretrained("Shrunoti09/llama2")
tokenizer.eos_token_id = model.config.eos_token_id
tokenizer.pad_token = tokenizer.eos_token
tokenizer.add_special_tokens({'pad_token': '<PAD>'})
if tokenizer.pad_token is None:
    tokenizer.pad_token = tokenizer.eos_token

app = Flask(__name__)
CORS(app,origins='*',resource={
    r"/*":{
        "origins":"*"
    }
})
# socketio = SocketIO(app,cors_allowed_origins='*')


def generate_words(section,message,file_name,tokens):

    message = str(message)
    # Generate words one by one
    text = f"<s>[INST] {message} [/INST]"

    inputs = tokenizer(text, return_tensors="pt")
    streamer = TextIteratorStreamer(tokenizer)
    # outputs = model.generate(input_ids=inputs["input_ids"].to("cuda"), streamer=streamer, max_new_tokens=400)


    generation_kwargs = dict(inputs, streamer=streamer, max_new_tokens=tokens, repetition_penalty = 1.1)
    thread = Thread(target=model.generate, kwargs=generation_kwargs)
    thread.start()

    paragraph = ""
    existing_sentences = set()
    generated_sentence = ""
    for new_text in streamer:
        # Ensure consistent sentence end detection
        generated_sentence += new_text
        if new_text in ".?!":  # Check for common sentence end punctuation

            # Perform case-insensitive and punctuation-trimmed sentence comparison
            clean_sentence = str(generated_sentence).replace("\"","")
            if len(existing_sentences) == 0:
                text_repl = f"<s>[INST] {message} [/INST]"
                clean_sentence = str(clean_sentence).replace(text_repl,"")
            if clean_sentence not in existing_sentences:
                existing_sentences.add(clean_sentence)
                paragraph += clean_sentence
                splitted_sentence = str(clean_sentence).split(' ')
                for word in splitted_sentence:            
                    yield word + " "
                    time.sleep(0.2)

            generated_sentence = ""

    torch.cuda.empty_cache()
    gc.collect()
    yield "<END> "



@app.route('/')
def home():
    return "Home Hellow"


@app.route('/stream_words',methods=['POST'])
def stream_words():
    body = request.json
    return Response(generate_words(body["section"],body["prompt"],body["file_name"],body["token"]), mimetype='text/plain')



@app.route('/paraphrase',methods=['POST'])
def paraphrase_sent():
    body = request.json

    API_URL = "https://api-inference.huggingface.co/models/tuner007/pegasus_paraphrase"
    headers = {"Authorization": "Bearer {}".format(HF_TOKEN)}

    def query(payload):
        response = requests.post(API_URL, headers=headers, json=payload)
        return response.json()
        
    output = query({
        "inputs": body["message"],
    })
    return output


@app.route('/donwnloadFile',methods=['POST'])
def download_file():
    body = request.json['data']
    
    chats = body["chats"]
    file_name = body["file_name"]+".docx"

    title = body["title"]
    type = body["type"]
    document = Document()
    if type == "Research Paper":
        document = Document(r"./conference-template-a4.docx")
        
    paragraphs = document.paragraphs

    if type == "Report":
        document.add_heading(title, 0)

    # document.save(file_name)
    for i in chats:
        if i["diagram"] == True:
            imagestr = i["output"]
            im = Image.open(BytesIO(b64decode(imagestr.split(',')[1])))
            imgName = body["file_name"]+".png"
            im.save(imgName)
            # Write binary data to a file

            if type == "Research Paper":
                ind = sections["ProposedMethod"]
                
                run = paragraphs[ind].add_run()
                run.add_picture(imgName,width=Inches(5))
            else:
                document.add_heading("Architecture", level=1)
                document.add_picture(imgName, width=Inches(10))

        else:
            if type == "Research Paper":
                ind = sections[i["section"]]
                paragraphs[ind] = paragraphs[ind].clear()
                paragraphs[ind].text = i["output"]
                paragraphs[ind].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph_format = paragraphs[ind].paragraph_format
            else:
                document.add_heading(i["section"], level=1)
                paragraph = document.add_paragraph(i["output"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.5)  

        document.save(file_name)


    # Send the PDF file as a response
    return send_file(file_name, as_attachment=False)




@app.route('/getFile',methods=['POST'])
def convert_and_get_pdf():
    body = request.json['data']

    title = body["title"]
    chats = body["chats"]
    file_name = body["file_name"]+".docx"
    type = body["type"]
    document = Document()
    if type == "Research Paper":
        document = Document(r"./conference-template-a4.docx")
        
    paragraphs = document.paragraphs

    if type == "Report":
        document.add_heading(title, 0)

    # document.save(file_name)
    for i in chats:
        if i["diagram"] == True:
            imagestr = i["output"]
            im = Image.open(BytesIO(b64decode(imagestr.split(',')[1])))
            imgName = body["file_name"]+".png"
            im.save(imgName)
            # Write binary data to a file
            if type == "Research Paper":
                ind = sections["ProposedMethod"]
                
                run = paragraphs[ind].add_run()
                run.add_picture(imgName,width=Inches(5))
            else:
                document.add_heading("Architecture", level=1)
                document.add_picture(imgName, width=Inches(10))
        else:       
            if type == "Research Paper":
                ind = sections[i["section"]]
                paragraphs[ind] = paragraphs[ind].clear()
                paragraphs[ind].text = i["output"]
                paragraphs[ind].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph_format = paragraphs[ind].paragraph_format
            else:
                document.add_heading(i["section"], level=1)
                paragraph = document.add_paragraph(i["output"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.5)

        document.save(file_name)

    pdf_file_path = body["file_name"]+".pdf"


    convertapi.convert('pdf', {
        'File': file_name
    }, from_format = 'docx').save_files(pdf_file_path)
    
    # Send the PDF file as a response
    return send_file(pdf_file_path, as_attachment=False)

if __name__ == '__main__':
    app.run(host='0.0.0.0',port=8000)
